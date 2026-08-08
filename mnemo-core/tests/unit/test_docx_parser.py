"""Unit tests for the DOCX parser."""

import io

import docx
import pytest
from mnemo.interfaces.errors import ContractValidationError
from mnemo.interfaces.parser_models import (
    RawHeadingBlock,
    RawImageBlock,
    RawListBlock,
    RawTableBlock,
    RawTextBlock,
)
from mnemo.interfaces.types import FileMetadata
from mnemo.parsers.docx import DOCXParser


def _create_test_docx() -> bytes:
    doc = docx.Document()
    doc.add_heading("Heading 1", level=1)
    doc.add_paragraph("Paragraph 1")
    doc.add_paragraph("List item 1", style="List Bullet")
    doc.add_paragraph("List item 2", style="List Bullet")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def parser() -> DOCXParser:
    return DOCXParser()


@pytest.fixture
def file_metadata() -> FileMetadata:
    return FileMetadata(
        content_hash="e3b0c44298fc1c149afbf4c8996fb92427ae41e4649b934ca495991b7852b855",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        size_bytes=100,
    )


def test_docx_parser_capabilities(parser: DOCXParser) -> None:
    caps = parser.capabilities()
    assert ".docx" in caps.supported_formats
    assert caps.supports_tables is True
    assert caps.supports_images is True


def test_docx_parser_empty_data(parser: DOCXParser, file_metadata: FileMetadata) -> None:
    with pytest.raises(ContractValidationError, match="Cannot parse empty DOCX"):
        parser.parse(b"", "test.docx", file_metadata)


def test_docx_parser_invalid_data(parser: DOCXParser, file_metadata: FileMetadata) -> None:
    with pytest.raises(ContractValidationError, match="Failed to open DOCX"):
        parser.parse(b"not a docx", "test.docx", file_metadata)


def test_docx_parser_extracts_blocks(parser: DOCXParser, file_metadata: FileMetadata) -> None:
    data = _create_test_docx()
    result = parser.parse(data, "test.docx", file_metadata)

    blocks = result.blocks
    assert len(blocks) == 4

    assert isinstance(blocks[0], RawHeadingBlock)
    assert blocks[0].text == "Heading 1"
    assert blocks[0].level == 1
    assert blocks[0].ordinal == 0

    assert isinstance(blocks[1], RawTextBlock)
    assert blocks[1].text == "Paragraph 1"
    assert blocks[1].ordinal == 1

    assert isinstance(blocks[2], RawListBlock)
    assert blocks[2].items == ("List item 1", "List item 2")
    assert blocks[2].ordinal == 2

    assert isinstance(blocks[3], RawTableBlock)
    assert blocks[3].rows == (("A1", "B1"), ("A2", "B2"))
    assert blocks[3].header_row_count == 1
    assert blocks[3].ordinal == 3


def test_docx_parser_extracts_images(parser: DOCXParser, file_metadata: FileMetadata) -> None:
    doc = docx.Document()
    doc.add_paragraph("Paragraph with image")

    # Create a valid dummy image (1x1 GIF)
    image_bytes = (
        b"GIF89a\x01\x00\x01\x00\x80\x00\x00\x00\x00\x00\xff\xff\xff"
        b"!\xf9\x04\x01\x00\x00\x00\x00,\x00\x00\x00\x00\x01\x00\x01"
        b"\x00\x00\x02\x01D\x00;"
    )
    image_stream = io.BytesIO(image_bytes)

    doc.add_picture(image_stream)

    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    result = parser.parse(data, "test.docx", file_metadata)

    # Text block, then Image block
    assert len(result.blocks) == 2

    img_block = result.blocks[1]
    assert isinstance(img_block, RawImageBlock)
    assert img_block.parser_local_id == "image-1"

    assert len(result.extracted_assets) == 1
    asset = result.extracted_assets[0]
    assert asset.parser_local_id == "image-1"
    assert asset.raw_bytes == image_bytes
    assert "image" in asset.mime_type
